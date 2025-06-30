import streamlit as st
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Mm
import os
import tempfile
import shutil

st.set_page_config(page_title="PDF转B5 Word（上下拼图居中）", layout="centered")
st.title("📄 PDF转B5 Word：每页拼接两张 PDF 页图（居中）")

uploaded_file = st.file_uploader("📤 上传 PDF 文件", type=["pdf"])

if uploaded_file:
    with st.spinner("正在处理 PDF..."):

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        pdf_path = os.path.join(temp_dir, "input.pdf")
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.read())

        # PDF转图片
        pdf = fitz.open(pdf_path)
        image_paths = []
        for i, page in enumerate(pdf):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 提高分辨率
            image_path = os.path.join(temp_dir, f"page_{i+1}.png")
            pix.save(image_path)
            image_paths.append(image_path)
        pdf.close()

        # 创建 Word 文档（B5尺寸竖版，176mm x 250mm）
        doc = Document()
        section = doc.sections[0]
        section.page_width = Mm(176)
        section.page_height = Mm(250)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)

        def insert_centered_image(image_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1  # 1=center
            img = Image.open(image_path)
            width_px, height_px = img.size
            img.close()
            dpi = 96
            width_in = width_px / dpi
            width_mm = width_in * 25.4
            max_width_mm = 156  # 页面宽度减去页边距
            scale = min(max_width_mm / width_mm, 1.0)
            doc_width_mm = width_mm * scale
            paragraph.add_run().add_picture(image_path, width=Mm(doc_width_mm))

        # 每页插入两张图片，上下居中排布
        for i in range(0, len(image_paths), 2):
            insert_centered_image(image_paths[i])
            doc.add_paragraph()
            if i + 1 < len(image_paths):
                insert_centered_image(image_paths[i + 1])
            if i + 2 < len(image_paths):
                doc.add_page_break()

        # 保存输出
        output_path = os.path.join(temp_dir, "b5_output.docx")
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.success("✅ 转换完成！点击下载：")
            st.download_button("📥 下载 B5 Word 文档", f, file_name="B5拼图输出.docx")

        shutil.rmtree(temp_dir)
